# -*- coding: utf-8 -*-
"""
Requirements (install first):
  - pip install PyMuPDF python-docx openpyxl pillow pytesseract hijridate pandas
Also install Tesseract OCR on your OS:
  - Windows: https://github.com/UB-Mannheim/tesseract/wiki
  - Linux: apt-get install tesseract-ocr
  - macOS (brew): brew install tesseract
Add language packs as needed (eng, urd, ara, hin, etc) to tessdata.
"""
import math
import os
import re
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import pandas as pd
from datetime import datetime, timedelta
from hijridate import Hijri
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from openpyxl import load_workbook
from openpyxl.styles import PatternFill, Alignment, Border, Side, Font

# -----------------------------
# 1) CONFIG: holidays/vacations
# -----------------------------
fixed_holidays = {
    "02-05": "Kashmir Day",
    "03-23": "Pakistan Day",
    "05-01": "Labour Day",
    "08-14": "Independence Day",
    "12-25": "Quaid-e-Azam Day / Christmas"
}
vacation_ranges = [("06-01", "08-10"), ("12-20", "12-31")]  # adjust as needed

# ---------------------------------------
# 2) Utilities: digit/script normalization
# ---------------------------------------
ARABIC_INDIC = str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789")       # Arabic-Indic
EXT_ARABIC_INDIC = str.maketrans("۰۱۲۳۴۵۶۷۸۹", "0123456789")   # Persian
DEVANAGARI = str.maketrans("०१२३४५६७८९", "0123456789")        # Hindi

def normalize_digits(text: str) -> str:
    """Convert Arabic-Indic/Persian/Devanagari digits to ASCII 0-9."""
    return text.translate(ARABIC_INDIC).translate(EXT_ARABIC_INDIC).translate(DEVANAGARI)

def clean_lines(text: str):
    """Basic cleanup and return list of non-empty trimmed lines without excessive repeats."""
    text = re.sub(r'\r', '\n', text)
    text = re.sub(r'\n+', '\n', text)
    text = normalize_digits(text)
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    # Remove frequent header/footer noise: drop lines that appear on too many pages by heuristic.
    # Simple heuristic: if a line repeats > 10 times, consider it boilerplate and drop.
    freq = {}
    for ln in lines:
        freq[ln] = freq.get(ln, 0) + 1
    return [ln for ln in lines if freq[ln] <= 10]

def is_skip_line(ln: str) -> bool:
    """Return True if line is a page number or header boilerplate."""
    if not ln or re.match(r'^\s*page\s+\d+\s*$', ln, re.IGNORECASE):
        return True
    return False

# ---------------------------------------------------

# 3) I/O: read any file (PDF, DOCX, image) + OCR text
# ---------------------------------------------------
def extract_text_any(path: str, ocr_langs: str = "eng+urd+ara+hin") -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext in (".pdf",):
        return extract_text_from_pdf(path, ocr_langs)
    elif ext in (".docx",):
        return extract_text_from_docx(path)
    elif ext in (".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"):
        return extract_text_from_image(path, ocr_langs)
    else:
        # Fallback: try PDF first; else OCR image; else read as plain text
        try:
            return extract_text_from_pdf(path, ocr_langs)
        except Exception:
            try:
                return extract_text_from_image(path, ocr_langs)
            except Exception:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()

def extract_text_from_pdf(path: str, ocr_langs: str) -> str:
    """
    Strategy:
      1) Try native text via PyMuPDF with page tags [Page X].
      2) If a page is mostly image/no text, OCR that page image with Tesseract.
    """
    doc = fitz.open(path)
    out = []
    for i, page in enumerate(doc, 1):
        t = page.get_text().strip()
        if t and len(t) > 20:
            out.append(f"[Page {i}]\n" + t)
        else:
            try:
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                t_ocr = pytesseract.image_to_string(img, lang=ocr_langs).strip()
                out.append(f"[Page {i}]\n" + t_ocr)
            except Exception:
                out.append(f"[Page {i}]")
    return "\n\n".join(out)


def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    # also read tables (sometimes books put content in tables)
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)

def extract_text_from_image(path: str, ocr_langs: str) -> str:
    img = Image.open(path)
    return pytesseract.image_to_string(img, lang=ocr_langs)

# ----------------------------------------------------------------
# 4) Detect content start (skip preface/TOC/credits) multilingual
# ----------------------------------------------------------------
# Multilingual markers to IGNORE (front matter)
SKIP_HINTS = [
    # English
    "preface", "foreword", "acknowledgement", "acknowledgment",
    "contents", "table of contents", "index", "glossary",
    "copyright", "publisher", "published by", "printing press",
    "printed by", "disclaimer", "isbn", "issn", "edition",
    "revised by", "edited by", "compiled by", "written by",
    "authors", "author", "reviewer", "proofreader", "printer",
    "all rights reserved", "note for teacher", "teachers' note",
    "teacher guide", "department of", "faculty of", "institute of",
    "university", "college", "chairman", "director", "principal",
    "headmaster", "headmistress", "professor", "dr ",
    # Specific name you mentioned
    "dr irshad ahmad",

    # Urdu
    "دیباچہ", "مقدمہ", "فہرست", "مشمولات", "حقوق", "اشاعت",
    "ڈیسکلیمر", "کتاب خانہ", "اشاعتی ادارہ", "پرنٹنگ پریس",
    "تصحیح", "مرتب", "مصنف", "معلم کے لئے نوٹ",
    "اساتذہ کے لئے نوٹ", "جامعہ", "یونیورسٹی", "کالج",
    "ڈاکٹر", "پروفیسر", "ڈاکٹر ارشاد احمد",

    # Arabic
    "مقدمة", "فهرس", "حقوق", "الناشر", "تمهيد", "إقرار",
    "تنبيه", "ملاحظات للمعلم", "الجامعة", "الكلية",
    "الدكتور", "الأستاذ",

    # Hindi
    "भूमिका", "प्राक्कथन", "सूची", "विषय सूची", "अनुक्रमणिका",
    "कॉपीराइट", "प्रकाशक", "अस्वीकरण", "आईएसबीएन",
    "लेखक", "संपादित", "संकलित", "परिशोधित",
    "विश्वविद्यालय", "महाविद्यालय", "डॉ"
]

# Multilingual headings for CHAPTER/UNIT/LESSON/SECTION
CHAPTER_TOKENS = [
    # English
    "unit", "chapter", "lesson", "section", "part", "module", "topic",
    # Arabic
    "الفصل", "الوحدة", "الدرس", "القسم", "الجزء",
    # Urdu
    "باب", "سبق", "یونٹ", "حصہ",
    # Hindi
    "अध्याय", "इकाई", "पाठ", "खंड"
]


def likely_heading(line: str) -> bool:
    """Check if a line looks like start of subject content by presence of known tokens or numbering."""
    low = line.lower()
    # token presence (any language)
    for tk in CHAPTER_TOKENS:
        if tk in low or tk in line:
            return True
    # numbered patterns like "1", "1.", "1.2", "1.2.3", "1) ...", "A) ..."
    if re.match(r"^([0-9]+(\.[0-9]+)*)[\s\).:-]", line):
        return True
    if re.match(r"^[A-Za-z]{1,3}[\)\.:-]\s+", line):  # A) , i) , a.
        return True
    # ALL CAPS big words can be false positive; keep conservative
    return False

def detect_content_start(lines):
    """
    Find the first line that looks like a *real* chapter/unit/lesson heading
    and is NOT part of known front matter hints.
    """
    def is_skip(line):
        low = line.lower()
        return any(h in low or h in line for h in SKIP_HINTS)

    for i, ln in enumerate(lines):
        if is_skip(ln):
            continue
        if likely_heading(ln):
            return i
    # fallback: start from line 0
    return 0

# -------------------------------------------------------
# 5) Multilingual structure extraction (no topic skipping)
# -------------------------------------------------------
def extract_structure_any(text: str):
    """
    Multilingual parsing that:
    - Detects chapter/unit/lesson/section headings (Urdu/Arabic/Hindi/English)
    - Detects numbered topics/subtopics
    - Never discards leftover lines: everything gets attached somewhere
    """
    lines = clean_lines(text)
    start_idx = detect_content_start(lines)
    lines = lines[start_idx:]
    lines = [normalize_digits(ln) for ln in lines]

    # token + number (supports Roman numerals too)
    chapter_re = re.compile(r"""
    ^(
       # English tokens
       (Unit|Chapter|Lesson|Section|Part|Module|Topic)\s*(\d+|[IVXLCDM]+)\b
      |# Arabic tokens
       (الفصل|الوحدة|الدرس|القسم|الجزء)\s*\d+\b
      |# Urdu tokens
       (باب|سبق|یونٹ|حصہ)\s*\d+\b
      |# Hindi tokens
       (अध्याय|इकाई|पाठ|खंड)\s*\d+\b
    )$
    """, re.IGNORECASE | re.VERBOSE)

    # number first, then token (e.g., "1. Chapter", "II - Unit")
    chapter_num_first = re.compile(r"""
    ^
     (\d+|[IVXLCDM]+)\s*[\)\.\:\-–]?\s*
     (Unit|Chapter|Lesson|Section|Part|Module|Topic|
      الفصل|الوحدة|الدرس|القسم|الجزء|
      باب|سبق|یونٹ|حصہ|
      अध्याय|इकाई|पाठ|खंड)\b
    """, re.IGNORECASE | re.VERBOSE)

    # Topic: hierarchical numbers like 1.2 or 2.3.4:
    topic_re = re.compile(r"^(\d+(\.\d+)+)[\s\.:–-]+(.+)?")

    # Subtopic: bullets like a), i), 1), etc. (neutral to language)
    subtopic_re = re.compile(r"^([A-Za-z0-9آ-یء]{1,3}[\)\.])\s+(.+)")

    structure = []
    current_ch = None
    current_tp = None

    for ln in lines:
        if is_skip_line(ln):
            continue


        # Chapter/Unit/Lesson/Section detect (both styles)
        if chapter_re.match(ln) or chapter_num_first.match(ln):
            current_ch = {'name': ln, 'topics': []}
            structure.append(current_ch)
            current_tp = None
            continue

        tp = topic_re.match(ln)
        if tp:
            if current_ch is None:
                current_ch = {'name': 'Unit 0', 'topics': []}
                structure.append(current_ch)
            current_tp = {'name': ln, 'subtopics': []}
            current_ch['topics'].append(current_tp)
            continue

        st = subtopic_re.match(ln)
        if st and current_tp:
            current_tp['subtopics'].append(ln)
            continue

        # If nothing matched but we have a current topic and the line is longish,
        # treat it as a continuation/subtopic (don’t lose content).
        if current_tp and len(ln) > 20:
            current_tp['subtopics'].append(ln)
            continue

        # If nothing matched at all and no chapter yet, create a bucket.
        if not current_ch and len(ln) > 20:
            current_ch = {'name': 'Unit 0', 'topics': []}
            structure.append(current_ch)
            current_tp = {'name': ln[:60], 'subtopics': []}
            current_ch['topics'].append(current_tp)

    # Ensure every chapter has at least one topic
    for ch in structure:
        if not ch['topics']:
            ch['topics'].append({'name': ch['name'], 'subtopics': []})

    return structure

def flatten_structure(structure):
    flat = []
    for ch in structure:
        ch_name = ch['name']
        for tp in ch['topics']:
            tp_name = tp['name']
            if tp['subtopics']:
                for st in tp['subtopics']:
                    flat.append((ch_name, tp_name, st))
            else:
                flat.append((ch_name, tp_name, ""))  # don't lose empty topics
    return flat
# --- helper: extract "1.2.3" style number + main title from a topic line ---
# --- helper: extract "1.2.3" style number + main title from a line ---
def parse_topic_num_title(tp: str):
    """
    Input: '1.2 States of Matter' or 'Unit 1 – Kinematics'
    Return: (num, title)
    """
    if not tp:
        return ("", "")
    # Try dotted numbers first (1.2, 2.3.4)
    m = re.match(r'^\s*(\d+(?:\.\d+)+)\s*[-–\.:]*\s*(.*)$', tp.strip())
    if m:
        num = m.group(1).strip()
        title = (m.group(2) or "").strip()
        return (num, title if title else tp.strip())
    # Try token + number (Unit 1, Chapter II)
    m2 = re.match(r'^\s*(Unit|Chapter|Lesson|Section|Part|Module|Topic|الفصل|الوحدة|الدرس|القسم|الجزء|باب|سبق|یونٹ|حصہ|अध्याय|इकाई|पाठ|खंड)\s*(\d+|[IVXLCDM]+)\s*[-–\.:]*\s*(.*)$', tp.strip(), flags=re.IGNORECASE)
    if m2:
        num = m2.group(2).strip()
        title = (m2.group(3) or "").strip()
        return (num, title if title else tp.strip())
    # Fallback: no number found
    return ("", tp.strip())

# ------------------------------
# 6) Holidays (same as your code)
# ------------------------------
def get_eid_holidays(year):
    # NOTE: this is an approximation; replace with an official calendar if needed.
    hijri_year = year - 579
    eid_ul_fitr = Hijri(hijri_year, 10, 1).to_gregorian()
    eid_ul_adha = Hijri(hijri_year, 12, 10).to_gregorian()
    return {
        eid_ul_fitr.strftime('%Y-%m-%d'): "Eid-ul-Fitr",
        (eid_ul_fitr + timedelta(days=1)).strftime('%Y-%m-%d'): "Eid-ul-Fitr Holiday",
        (eid_ul_fitr + timedelta(days=2)).strftime('%Y-%m-%d'): "Eid-ul-Fitr Holiday",
        eid_ul_adha.strftime('%Y-%m-%d'): "Eid-ul-Adha",
        (eid_ul_adha + timedelta(days=1)).strftime('%Y-%m-%d'): "Eid-ul-Adha Holiday",
        (eid_ul_adha + timedelta(days=2)).strftime('%Y-%m-%d'): "Eid-ul-Adha Holiday"
    }

def is_holiday(date):
    d_str = date.strftime('%Y-%m-%d')
    mmdd = date.strftime('%m-%d')
    if mmdd in fixed_holidays:
        return f"Holiday: {fixed_holidays[mmdd]}"
    for start, end in vacation_ranges:
        start_d = datetime.strptime(f"{date.year}-{start}", "%Y-%m-%d")
        end_d   = datetime.strptime(f"{date.year}-{end}",   "%Y-%m-%d")
        if start_d <= date <= end_d:
            return "Vacation"
    eid_days = get_eid_holidays(date.year)
    if d_str in eid_days:
        return f"Holiday: {eid_days[d_str]}"
    return None

# -------------------------------------------------
# 7) Schedule exactly 9 months from start (teaching)
# -------------------------------------------------
def add_months(dt, months):
    # Simple month add without external deps
    year = dt.year + (dt.month - 1 + months) // 12
    month = (dt.month - 1 + months) % 12 + 1
    day = min(dt.day, [31,
        29 if year % 4 == 0 and (year % 100 != 0 or year % 400 == 0) else 28,
        31,30,31,30,31,31,30,31,30,31][month-1])
    return datetime(year, month, day)

def generate_curriculum_9_months(start_date, structure, subject, grade):
    """
    Months are RELATIVE to session start:
      Months 1–9  -> teaching (finish exactly by end of month 9)
      Month 10    -> chapter/unit-wise tests
      Month 11    -> Half/Full book tests (alternate)
      Month 12    -> Final Exam mid-month; other days revision
    """
    def rel_month(d):
        # 1-based month index relative to session start
        return (d.year - start_date.year) * 12 + (d.month - start_date.month) + 1

    data = []
    week = 1

    # We need 12 months total (teaching 1–9, tests 10–12)
    end_date = add_months(start_date, 12) - timedelta(days=1)

    # Build full date range
    all_days = []
    temp = start_date
    while temp <= end_date:
        all_days.append(temp)
        temp += timedelta(days=1)

    # ---------- Teaching days plan (only months 1–9) ----------
    teaching_days = []
    for d in all_days:
        if rel_month(d) <= 9 and is_holiday(d) is None and d.strftime('%A') not in ['Saturday', 'Sunday']:
            teaching_days.append(d)

    # Flatten syllabus in strict order (no skipping)
    flat = flatten_structure(structure)  # list of (chapter, topic, subtopic)
    total_items = len(flat)
    if total_items == 0:
        raise ValueError("No syllabus items parsed. Check parsing or input file.")

    # Apportion topics over teaching days so we finish exactly by end of month 9
    remaining_items = total_items
    remaining_workdays = len(teaching_days)
    items_plan = {}
    for wd in teaching_days:
        items_today = max(1, math.ceil(remaining_items / remaining_workdays))
        items_today = min(items_today, remaining_items)
        items_plan[wd] = items_today
        remaining_items -= items_today
        remaining_workdays -= 1

    idx = 0  # pointer into flat syllabus

    # Anchors for month 10/11 indexing (relative)
    month10_start = add_months(start_date, 9)   # start of relative month 10
    month11_start = add_months(start_date, 10)  # start of relative month 11

    # Chapter list for month 10 tests (skip "Unit 0" if present)
    chapters_for_tests = [c for c in structure if not re.match(r'^\s*unit\s*0\b', c['name'], re.I)]
    if not chapters_for_tests:
        chapters_for_tests = structure[:]  # fallback

    # ---------- Build the calendar rows ----------
    for d in all_days:
        date_str = d.strftime('%Y-%m-%d')
        day_name = d.strftime('%A')
        month_name = d.strftime('%B')
        highlight = ""
        topic = subtopic = activity = homework = quiz = "-"

        # Holidays/Weekends
        holiday_tag = is_holiday(d)
        if holiday_tag:
            topic = holiday_tag if holiday_tag != "Vacation" else "Vacation"
            subtopic = "No academic activity"
            highlight = "red"
        elif day_name in ['Saturday', 'Sunday']:
            topic = f"{day_name} Off"
            subtopic = "No academic activity"
            highlight = "green"
        else:
            # Which relative month is this date?
            rm = rel_month(d)

            if rm <= 9:
                # Teaching (exactly as your original logic)
                teach_count = items_plan.get(d, 0)
                if teach_count > 0:
                    first_num = first_title = ""
                    last_num = last_title = ""
                    first_unit_num = first_unit_title = ""
                    last_unit_num = last_unit_title = ""
                    taught_any = False

                    for _ in range(teach_count):
                        ch, tp, st = flat[idx]  # (chapter_line, topic_line, subtopic_line)
                        u_num, u_title = parse_topic_num_title(ch)
                        t_num, t_title = parse_topic_num_title(tp)

                        if not taught_any and (t_num or t_title):
                            first_num, first_title = t_num, t_title
                            first_unit_num, first_unit_title = u_num, u_title
                            taught_any = True
                        if (t_num or t_title):
                            last_num, last_title = t_num, t_title
                            last_unit_num, last_unit_title = u_num, u_title

                        idx += 1

                    # Build unit label
                    def unit_label(num, title):
                        if num and title:
                            return f"Unit {num} – {title}"
                        if num:
                            return f"Unit {num}"
                        return title or ""

                    # Unit range
                    if (first_unit_num, first_unit_title) == (last_unit_num, last_unit_title):
                        unit_str = unit_label(first_unit_num, first_unit_title)
                    else:
                        left_u = unit_label(first_unit_num, first_unit_title) or "Unit start"
                        right_u = unit_label(last_unit_num, last_unit_title) or "Unit end"
                        unit_str = f"{left_u} → {right_u}"

                    # Topic range
                    if (first_num, first_title) == (last_num, last_title) or not (last_num or last_title):
                        topic_range = f"{first_num} {first_title}".strip() if first_num else (first_title or "Topic")
                    else:
                        left = (f"{first_num} {first_title}").strip() if (first_num or first_title) else "Start"
                        right = (f"{last_num} {last_title}").strip() if (last_num or last_title) else "End"
                        topic_range = f"{left} → {right}"

                    topic = f"{unit_str}: {topic_range}" if unit_str else topic_range
                    subtopic = f"{teach_count} item(s)"
                    activity = "Lecture / Activity"
                    homework = "Home task"
                else:
                    topic = "Buffer / Reinforcement"
                    subtopic = "Light workload"
                    activity = "Reinforcement"

            elif rm == 10:
                # Chapter/Unit-wise tests (working days only)
                if chapters_for_tests:
                    day_index = (d - month10_start).days
                    chap_idx = day_index % len(chapters_for_tests)
                    chap_name = chapters_for_tests[chap_idx]['name']
                else:
                    chap_name = "Chapter"
                topic = f"{chap_name} Test"
                subtopic = f"Test/Revision: {chap_name}"
                activity = "Chapter Test / Oral / Written"

            elif rm == 11:
                # Alternate Half / Full book tests (working days only)
                day_index = (d - month11_start).days
                if day_index % 2 == 0:
                    topic = "Half Book Test"
                    subtopic = "First Half of Book"
                else:
                    topic = "Full Book Test"
                    subtopic = "Entire Book"
                activity = "Written Test"

            elif rm == 12:
                # Final Exam in the middle week; rest revision
                if 12 <= d.day <= 18:
                    topic = "Final Exam"
                    subtopic = "Board Style Exam"
                    activity = "Final Written Exam"
                else:
                    topic = "Revision / Project Work"
                    subtopic = "Catch-up and Revision"
                    activity = "Remedial / Project"

            # Friday quiz across all months
            if day_name == "Friday":
                quiz = "Quiz"

        data.append({
            'Date': date_str,
            'Day': day_name,
            'Week': week,
            'Month': month_name,
            'Subject': subject,
            'Grade': grade,
            'Topic': topic,
            'Subtopic': subtopic,
            'Activity': activity,
            'Homework': homework,
            'Quiz/Test': quiz,
            'Highlight': highlight
        })

        if day_name == 'Friday':
            week += 1

    return pd.DataFrame(data)

def generate_curriculum_custom(start_date, structure, subject, grade, duration_val=9, duration_unit='months'):
    """
    Generate study curriculum for ANY user-selected timeframe:
    duration_unit: 'days' | 'weeks' | 'months'
    duration_val: integer count (e.g. 14, 30, 6, 9, 12, etc.)
    """
    try:
        duration_val = int(duration_val)
    except Exception:
        duration_val = 9

    unit = (duration_unit or 'months').lower().strip()

    if unit == 'days':
        total_days = max(1, duration_val)
        end_date = start_date + timedelta(days=total_days - 1)
    elif unit == 'weeks':
        total_days = max(1, duration_val * 7)
        end_date = start_date + timedelta(days=total_days - 1)
    else:  # months
        end_date = add_months(start_date, max(1, duration_val)) - timedelta(days=1)

    all_days = []
    temp = start_date
    while temp <= end_date:
        all_days.append(temp)
        temp += timedelta(days=1)

    flat = flatten_structure(structure)
    total_items = len(flat)
    if total_items == 0:
        raise ValueError("No syllabus items parsed from input file.")

    if unit == 'days' or (unit == 'weeks' and duration_val <= 2):
        teaching_days = all_days[:]
    else:
        teaching_days = [d for d in all_days if is_holiday(d) is None and d.strftime('%A') not in ['Saturday', 'Sunday']]
        if not teaching_days:
            teaching_days = all_days[:]

    remaining_items = total_items
    remaining_workdays = len(teaching_days)
    items_plan = {}
    for wd in teaching_days:
        items_today = max(1, math.ceil(remaining_items / remaining_workdays))
        items_today = min(items_today, remaining_items)
        items_plan[wd] = items_today
        remaining_items -= items_today
        remaining_workdays -= 1

    data = []
    week = 1
    idx = 0

    for d_idx, d in enumerate(all_days, 1):
        date_str = d.strftime('%Y-%m-%d')
        day_name = d.strftime('%A')
        month_name = d.strftime('%B')
        highlight = ""

        holiday_tag = is_holiday(d)
        if holiday_tag and unit != 'days':
            topic = holiday_tag
            subtopic = "No academic activity"
            activity = "Rest / Holiday"
            homework = "-"
            quiz = "-"
            highlight = "red"
        elif day_name in ['Saturday', 'Sunday'] and unit != 'days' and (unit == 'months' or duration_val > 2):
            topic = f"{day_name} Off"
            subtopic = "Self-Study & Review"
            activity = "Review Week's Work"
            homework = "Review pending topics"
            quiz = "-"
            highlight = "green"
        else:
            teach_count = items_plan.get(d, 0)
            if teach_count > 0 and idx < total_items:
                ch_set, tp_set, st_set = [], [], []
                for _ in range(teach_count):
                    if idx < total_items:
                        ch, tp, st = flat[idx]
                        if ch and ch not in ch_set: ch_set.append(ch)
                        if tp and tp not in tp_set: tp_set.append(tp)
                        if st and st not in st_set: st_set.append(st)
                        idx += 1

                u_title = ch_set[0] if ch_set else "Unit"
                topic = " | ".join(tp_set[:2]) if tp_set else "Core Concept Breakdown"
                subtopic = " | ".join(st_set[:2]) if st_set else f"Detailed study of {u_title}"
                activity = "Interactive Lecture & Guided Reading"
                homework = "Solve practice problems & textbook exercises"
                quiz = "Quiz" if day_name == 'Friday' or unit == 'days' else "-"
            else:
                topic = "Consolidation & Final Revision"
                subtopic = "Review previous chapters & practice questions"
                activity = "Self-Assessment / Mock Exam"
                homework = "Final review notes"
                quiz = "Comprehensive Assessment"

        data.append({
            'Date': date_str,
            'Day': day_name,
            'Week': week,
            'Month': month_name,
            'Subject': subject,
            'Grade': grade,
            'Topic': topic,
            'Subtopic': subtopic,
            'Activity': activity,
            'Homework': homework,
            'Quiz/Test': quiz,
            'Highlight': highlight
        })

        if day_name == 'Friday':
            week += 1

    return pd.DataFrame(data)

# -----------------------------------------
# 8) Exporters (Word & Excel) unchanged-ish
# -----------------------------------------
def export_to_word(df, filename):
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    doc.add_heading(f"{df['Subject'][0]} - Grade {df['Grade'][0]} Daily Curriculum", 0)

    for month in df['Month'].unique():
        doc.add_page_break()
        doc.add_heading(month, level=1)
        table = doc.add_table(rows=1, cols=len(df.columns)-1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        borders = OxmlElement('w:tblBorders')
        for edge in ['top','left','bottom','right','insideH','insideV']:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'),'single')
            el.set(qn('w:sz'),'6')
            el.set(qn('w:color'),'000000')
            borders.append(el)
        table._tbl.tblPr.append(borders)
        headers = df.columns[:-1]  # drop Highlight column in Word
        for i, col in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = col
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            run.bold = True
        for _, row in df[df['Month']==month].iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(headers):
                cells[i].text = str(row[col])
                p = cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p.runs:
                    p.runs[0].font.size = Pt(9)
            # shade the entire row
            if row['Highlight'] in ('red', 'green'):
                fill = 'FF9999' if row['Highlight']=='red' else '92D050'
                for c in cells:
                    shade = OxmlElement('w:shd')
                    shade.set(qn('w:fill'), fill)
                    c._tc.get_or_add_tcPr().append(shade)
    doc.save(filename)

def highlight_excel(excel_path, df):
    wb = load_workbook(excel_path)
    ws = wb.active
    red = PatternFill(start_color='FF9999', end_color='FF9999', fill_type='solid')
    green = PatternFill(start_color='92D050', end_color='92D050', fill_type='solid')
    thin = Border(left=Side(style='thin'), right=Side(style='thin'),
                  top=Side(style='thin'), bottom=Side(style='thin'))
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)

    # header
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.border = thin
        cell.alignment = center

    # body
    for i, row in enumerate(ws.iter_rows(min_row=2), start=0):
        tag = df.iloc[i]['Highlight']
        for cell in row:
            cell.border = thin
            cell.alignment = center
            if tag == 'red':
                cell.fill = red
            elif tag == 'green':
                cell.fill = green
    wb.save(excel_path)

# -----------------------
# 9) CLI main orchestrator
# -----------------------
def main():
    input_path = input("Book file (PDF/DOCX/IMG): ").strip()
    subject = input("Subject: ").strip()
    grade = input("Grade/Class: ").strip()
    start_date_str = input("Session Start Date (YYYY-MM-DD): ").strip()
    # OCR languages string for Tesseract: e.g., "eng+urd+ara+hin"
    ocr_langs = input("OCR Languages (e.g., eng+urd+ara+hin) [default: eng+urd+ara+hin]: ").strip() or "eng+urd+ara+hin"

    start_date = datetime.strptime(start_date_str, "%Y-%m-%d")

    print("Extracting content...")
    raw_text = extract_text_any(input_path, ocr_langs=ocr_langs)

    print("Parsing structure (multilingual)...")
    structure = extract_structure_any(raw_text)

    # If structure ended empty (extreme case), make one topic with the whole content
    if not structure:
        structure = [{'name': 'Unit 0', 'topics': [{'name': 'Content', 'subtopics': [raw_text]}]}]

    print("Chapters/Units found:", [c['name'] for c in structure])

    print("Building 9-month schedule...")
    df = generate_curriculum_9_months(start_date, structure, subject, grade)

    base = f"{subject}_Grade{grade}_Curriculum_{start_date.strftime('%Y%m%d')}_9M"
    excel_file = f"{base}.xlsx"
    word_file = f"{base}.docx"

    df.to_excel(excel_file, index=False)
    highlight_excel(excel_file, df)
    export_to_word(df, word_file)

    print(f"✅ Saved:\n📊 {excel_file}\n📝 {word_file}")

if __name__ == "__main__":
    main()  